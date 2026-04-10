import importlib
import logging
import re
from pathlib import Path
from typing import Iterable, List

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from app.core.config import get_settings

logger = logging.getLogger(__name__)
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.*)$")


class UnsupportedFileTypeError(ValueError):
    pass


def load_documents(file_path: str, metadata: dict) -> List[Document]:
    suffix = Path(file_path).suffix.lower()
    if suffix == '.txt':
        return _load_text(file_path, metadata)
    if suffix == '.md':
        return _load_markdown(file_path, metadata)
    if suffix == '.pdf':
        return _load_pdf(file_path, metadata)
    if suffix == '.docx':
        return _load_docx(file_path, metadata)
    raise UnsupportedFileTypeError(f'Unsupported file type: {suffix}')


def _load_text(file_path: str, metadata: dict) -> List[Document]:
    with open(file_path, 'r', encoding='utf-8-sig') as file:
        content = file.read().strip()
    meta = dict(metadata)
    meta['location'] = 'full-text'
    return [Document(page_content=content, metadata=meta)]


def _load_markdown(file_path: str, metadata: dict) -> List[Document]:
    with open(file_path, 'r', encoding='utf-8-sig') as file:
        content = file.read().strip()
    return _documents_from_markdown(content, metadata=metadata, page_number=None, fallback_location='markdown')


def _load_pdf(file_path: str, metadata: dict) -> List[Document]:
    settings = get_settings()
    preferred_extractor = settings.pdf_extractor.lower()
    if settings.use_pymupdf4llm and preferred_extractor in {'auto', 'pymupdf4llm'}:
        structured = _load_pdf_with_pymupdf4llm(file_path, metadata)
        if structured:
            return structured
        if preferred_extractor == 'pymupdf4llm':
            raise ValueError('PyMuPDF4LLM extraction was requested but is not available.')

    return _load_pdf_with_pypdf(file_path, metadata)


def _load_pdf_with_pymupdf4llm(file_path: str, metadata: dict) -> List[Document]:
    module = _load_optional_module('pymupdf4llm')
    if module is None:
        return []

    docs: list[Document] = []
    try:
        page_chunks = module.to_markdown(file_path, page_chunks=True)
    except Exception as exc:
        logger.warning('PyMuPDF4LLM extraction failed for %s: %s', file_path, exc)
        return []

    if isinstance(page_chunks, str):
        return _documents_from_markdown(page_chunks, metadata=metadata, page_number=None, fallback_location='pdf')

    for page_index, item in enumerate(page_chunks or [], start=1):
        markdown = (item.get('text') or item.get('markdown') or '').strip()
        if not markdown:
            continue
        page_meta = item.get('metadata') or {}
        page_number = page_meta.get('page') or page_meta.get('page_number') or item.get('page') or page_index
        docs.extend(_documents_from_markdown(markdown, metadata=metadata, page_number=page_number, fallback_location='pdf'))
    return docs


def _load_pdf_with_pypdf(file_path: str, metadata: dict) -> List[Document]:
    reader = PdfReader(file_path)
    docs = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ''
        if not text.strip():
            continue
        meta = dict(metadata)
        meta['location'] = f'page-{index}'
        meta['page_number'] = index
        meta['citation_label'] = _build_citation_label(meta)
        docs.append(Document(page_content=text, metadata=meta))
    return docs


def _load_docx(file_path: str, metadata: dict) -> List[Document]:
    document = DocxDocument(file_path)
    content = '\n'.join([paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()])
    return _documents_from_markdown(content, metadata=metadata, page_number=None, fallback_location='docx')


def _documents_from_markdown(markdown: str, metadata: dict, page_number: int | None, fallback_location: str) -> List[Document]:
    sections = _split_markdown_sections(markdown)
    docs = []
    for index, section in enumerate(sections, start=1):
        text = section['content'].strip()
        if not text:
            continue
        meta = dict(metadata)
        meta['page_number'] = page_number
        meta['chapter_title'] = section['chapter_title']
        meta['section_title'] = section['section_title']
        meta['location'] = _build_location(page_number, section['section_title'], fallback_location, index)
        meta['citation_label'] = _build_citation_label(meta)
        docs.append(Document(page_content=text, metadata=meta))
    return docs


def _split_markdown_sections(markdown: str) -> list[dict]:
    chapter_title = ''
    section_title = ''
    buffer: list[str] = []
    sections: list[dict] = []

    def flush_buffer() -> None:
        content = '\n'.join(buffer).strip()
        if content:
            sections.append(
                {
                    'chapter_title': chapter_title,
                    'section_title': section_title or chapter_title,
                    'content': content,
                }
            )
        buffer.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        match = HEADING_PATTERN.match(line.strip())
        if match:
            flush_buffer()
            level = len(match.group(1))
            heading = match.group(2).strip()
            if level <= 2:
                chapter_title = heading
                section_title = heading
            else:
                section_title = heading
            continue
        buffer.append(line)

    flush_buffer()
    if sections:
        return sections
    return [{'chapter_title': '', 'section_title': '', 'content': markdown}]


def _build_location(page_number: int | None, section_title: str, fallback_location: str, index: int) -> str:
    parts = []
    if page_number:
        parts.append(f'page-{page_number}')
    if section_title:
        parts.append(section_title)
    if not parts:
        parts.append(f'{fallback_location}-{index}')
    return ' / '.join(parts)


def _build_citation_label(metadata: dict) -> str:
    parts = []
    if metadata.get('book_title'):
        parts.append(metadata['book_title'])
    if metadata.get('page_number'):
        parts.append(f"p.{metadata['page_number']}")
    if metadata.get('section_title'):
        parts.append(metadata['section_title'])
    return ' | '.join(parts)


def _load_optional_module(name: str):
    try:
        return importlib.import_module(name)
    except ImportError:
        return None
